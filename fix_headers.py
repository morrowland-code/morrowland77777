from docx import Document
import json
import re

# --------------------------------------------------
# 1. Load the 243-code → archetype-name mapping
# --------------------------------------------------
with open("archetypes_full.json", "r", encoding="utf-8") as f:
    ARCHETYPES = json.load(f)

EXPECTED_CODES = set(ARCHETYPES.keys())

# --------------------------------------------------
# 2. Header detector (very forgiving)
# --------------------------------------------------
HEADER_RE = re.compile(
    r"(?i)"                                   # case-insensitive
    r"openness:\s*(low|medium|high)\s*\|\s*"
    r"conscientiousness:\s*(low|medium|high)\s*\|\s*"
    r"extraversion:\s*(low|medium|high)\s*\|\s*"
    r"agreeableness:\s*(low|medium|high)\s*\|\s*"
    r"neuroticism:\s*(low|medium|high)"       # allow extra junk after this
)

def normalize_cap(word: str) -> str:
    """'low' -> 'Low', 'medium' -> 'Medium'."""
    w = word.strip().lower()
    if w == "low":
        return "Low"
    if w == "medium":
        return "Medium"
    if w == "high":
        return "High"
    return word

def main():
    doc_name = "morrowland 243.docx"
    print(f"📄 Opening {doc_name} ...")
    doc = Document(doc_name)

    seen_codes = set()
    fixed_count = 0
    suspicious_headers = []

    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        if "Openness" not in text:
            continue

        m = HEADER_RE.search(text)
        if not m:
            # Looks like it's *trying* to be a header but doesn't match pattern
            suspicious_headers.append((i, text))
            continue

        # Extract trait levels
        O, C, E, A, N = [normalize_cap(x) for x in m.groups()]
        code = f"{O}-{C}-{E}-{A}-{N}"
        seen_codes.add(code)

        # Look up the archetype name
        name = ARCHETYPES.get(code, f"UNKNOWN_{code}")

        # ✅ Rewrite header line to the clean, exact format
        new_header = (
            f"Openness: {O} | Conscientiousness: {C} | "
            f"Extraversion: {E} | Agreeableness: {A} | Neuroticism: {N}"
        )
        p.text = new_header

        # ✅ Make sure the *next* paragraph is the Archetype line
        if i + 1 < len(doc.paragraphs):
            next_p = doc.paragraphs[i + 1]
            next_p.text = f"Archetype: {name}"
        else:
            # If for some reason there is no next paragraph, create one
            new_p = doc.add_paragraph()
            new_p.text = f"Archetype: {name}"

        fixed_count += 1

    # --------------------------------------------------
    # 3. Save fixed document
    # --------------------------------------------------
    out_name = "morrowland 243_fixed.docx"
    doc.save(out_name)

    print("✅ Done rewriting headers.")
    print(f"   ✔ Normalized {fixed_count} header blocks.")
    print(f"   ✔ Unique codes found in DOCX: {len(seen_codes)}")

    # Missing codes vs 243 expected
    missing_codes = sorted(list(EXPECTED_CODES - seen_codes))
    if missing_codes:
        print("\n⚠ Some codes from archetypes_full.json were NOT seen in the document:")
        print(f"   Missing count: {len(missing_codes)}")
        for code in missing_codes:
            print(f"    - {code}: {ARCHETYPES[code]}")
    else:
        print("\n🎉 All 243 codes were present in the document!")

    # Show suspicious "Openness" lines that did NOT match the pattern
    if suspicious_headers:
        print("\n⚠ These paragraphs contain 'Openness' but didn't match the pattern.")
        print("   Fix these in Word so they look like:")
        print('   "Openness: High | Conscientiousness: Medium | Extraversion: Medium | Agreeableness: Medium | Neuroticism: Medium"')
        print("   Then run this script again.\n")
        for idx, line in suspicious_headers:
            print(f"   Paragraph {idx}: {line}")

if __name__ == "__main__":
    main()