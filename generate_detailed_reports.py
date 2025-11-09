from docx import Document
import json, os

# load archetypes
with open("archetypes_full.json", "r", encoding="utf-8") as f:
    data = json.load(f)

template_sections = [
    "Overview", "Deep Trait Breakdown", "Decision-Making Style",
    "Cognitive / Intelligence Style", "Aesthetic & Style",
    "Symbolic Expression", "Hobbies & Lifestyle",
    "Art & Entertainment", "Relationships & Attraction",
    "Compatibility", "Social & Communication",
    "Career & Purpose", "Gifts & Preferences",
    "Ethical Influence", "Manipulation Defense",
    "Growth Plan", "Subtype", "Summary", "Quote"
]

def generate_report(name, code):
    doc = Document()
    doc.add_heading(f"{name} — {code}", 0)
    doc.add_paragraph(f"Detailed Personality Archetype Report for {name}")

    O, C, E, A, N = code.split("-")

    for section in template_sections:
        doc.add_heading(section, level=1)
        # placeholder text you can expand later
        doc.add_paragraph(f"This section describes how a person with "
                          f"Openness {O}, Conscientiousness {C}, "
                          f"Extraversion {E}, Agreeableness {A}, "
                          f"and Neuroticism {N} behaves in {section.lower()}.")

    filename = f"reports/{name.replace(' ', '_')}.docx"
    os.makedirs("reports", exist_ok=True)
    doc.save(filename)
    print(f"✓ Saved {filename}")

# generate all
for code, name in data.items():
    generate_report(name, code)

print("\nAll detailed reports generated in the /reports folder.")