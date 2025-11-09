# build_full_docx_rich.py
# Creates: Big_Five_Archetypes_243_FULL.docx
# Uses your existing mapping (archetypes_full.json or archetypes.json or the DOCX you built earlier)
# and writes 22 sections × 3 paragraphs per archetype in Tone A.

import os, re, json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DOC = "Big_Five_Archetypes_243_FULL.docx"

SECTION_TITLES = [
    "Archetype Identity","Core Personality Summary","Deep Trait Analysis","Internal World",
    "Cognitive & Intelligence Profile","Decision-Making Style","Aesthetic & Style","Symbolism & Tattoos",
    "Motivation & Purpose","Communication & Social Energy","Relationship Blueprint","Compatibility",
    "Attraction & Seduction Style","Gifts, Taste & Interests","Career, Power & Productivity",
    "Influence, Persuasion & Manipulation Profile","Defense & Shadow","Growth Path & Transformation",
    "Art, Entertainment & Culture Preferences","Symbolic Archetype Summary","MBTI Mirror","Final Quote"
]

# ---------- load your code->name map ----------
def load_map():
    # prefer a simple dict like {"Low-Low-...":"Aquashine", ...}
    if os.path.exists("archetypes_full.json"):
        with open("archetypes_full.json","r",encoding="utf-8") as f:
            d = json.load(f)
            # either {code:name} OR {name:{traits:"Low | ..."}}; normalize:
            if all("-" in k for k in d.keys()):
                return d
            # else build from nested
            m = {}
            for name,info in d.items():
                if isinstance(info,dict) and "traits" in info:
                    code = info["traits"].replace(" ","").replace("|","-")
                    m[code] = name
            if m: return m

    if os.path.exists("archetypes.json"):
        with open("archetypes.json","r",encoding="utf-8") as f:
            return json.load(f)

    # last resort: parse from docx with lines like
    # "Openness: Low | ... — Archetype: Aquashine"
    docx_source = "Big_Five_Archetypes_243 (1).docx"
    if os.path.exists(docx_source):
        doc = Document(docx_source)
        pat = re.compile(
            r"Openness:\s*(Low|Medium|High)\s*\|\s*"
            r"Conscientiousness:\s*(Low|Medium|High)\s*\|\s*"
            r"Extraversion:\s*(Low|Medium|High)\s*\|\s*"
            r"Agreeableness:\s*(Low|Medium|High)\s*\|\s*"
            r"Neuroticism:\s*(Low|Medium|High)\s*—\s*Archetype:\s*(.+)\s*$"
        )
        m = {}
        for p in doc.paragraphs:
            t = p.text.strip()
            mobj = pat.search(t)
            if mobj:
                O,C,E,A,N,name = mobj.groups()
                code = f"{O}-{C}-{E}-{A}-{N}"
                m[code] = name.strip()
        if m: 
            with open("archetypes_full.json","w",encoding="utf-8") as f:
                json.dump(m,f,ensure_ascii=False,indent=2)
            return m

    # minimal fallback so script still runs
    return {
        "Low-Low-Low-Low-Low": "Aquashine",
        "High-High-High-High-High": "Emberheart"
    }

def split5(code):
    parts = code.split("-")
    if len(parts)!=5: parts = ["Medium"]*5
    return parts

def trait_line(level, low, med, high):
    return {"Low":low,"Medium":med,"High":high}[level]

def three_paras(name, title, O,C,E,A,N):
    # short “Tone A” generator per section; 3 distinct paragraphs driven by traits
    open_p  = trait_line(O,
        "prefers known textures and trusted rituals",
        "mixes curiosity with practicality",
        "seeks novel textures, symbolism, and creative edges"
    )
    cons_p  = trait_line(C,
        "resists rigid scaffolding and moves by inner pacing",
        "keeps a workable plan with room to breathe",
        "thrives with structure, preparation, and clean follow-through"
    )
    extra_p = trait_line(E,
        "recharges in solitude and close circles",
        "shifts gracefully between company and quiet",
        "comes alive in motion, collaboration, and buzz"
    )
    agree_p = trait_line(A,
        "speaks plainly and guards boundaries",
        "weighs harmony with self-respect",
        "leads with warmth, accommodation, and care"
    )
    neuro_p = trait_line(N,
        "stays steady and difficult to perturb",
        "feels deeply yet recovers with reflection",
        "meets life intensely and benefits from grounding rituals"
    )

    if title == "Archetype Identity":
        return [
            f"{name} carries a five-tone signature—Openness {O}, Conscientiousness {C}, Extraversion {E}, Agreeableness {A}, Neuroticism {N}.",
            f"They {open_p}; they {cons_p}; they {extra_p}. In connection this type {agree_p}, while emotionally this person {neuro_p}.",
            "What people sense is coherence: values quietly steering choices, presence tuned to the right pace."
        ]
    if title == "Core Personality Summary":
        return [
            "At the core: equilibrium between instinct and intention—meaning over momentum.",
            "Strengths unfold across the long arc: fidelity, discernment, emotional memory, patience that turns into leverage.",
            "Even when contexts shift, integrity remains the through-line; reliability becomes quiet power."
        ]
    if title == "Deep Trait Analysis":
        return [
            f"Openness ({O}): they {open_p}.",
            f"Conscientiousness ({C}): they {cons_p}. Extraversion ({E}): they {extra_p}.",
            f"Agreeableness ({A}): they {agree_p}. Neuroticism ({N}): they {neuro_p}."
        ]
    if title == "Internal World":
        return [
            "Inner life moves in rhythms; thoughts circle, refine, and arrive whole.",
            "Emotion is metabolized by reflection and pattern—they don’t rush meaning; they cultivate it.",
            "Withdrawal is not disappearance but recentering—the compass resets, then they return."
        ]
    if title == "Cognitive & Intelligence Profile":
        return [
            "Intelligence shows as observation, synthesis, and tone-reading more than speed.",
            "They learn best from story and example; give them raw material and they’ll shape coherence.",
            "Insights arrive measured but sticky—what they say at minute sixty changes the room."
        ]
    if title == "Decision-Making Style":
        return [
            "They triangulate values, evidence, and felt sense—logic must be livable.",
            "Under pressure they slow the clock, buying clarity with time instead of trading it for urgency.",
            "Once aligned, they commit; applause never outranks accuracy."
        ]
    if title == "Aesthetic & Style":
        return [
            "They choose coherence over novelty: repeatable palettes, tactile comfort, meaningful objects.",
            "Clothing serves the day’s story—quietly expressive, never performative.",
            "Their spaces breathe: soft light, negative space, a few anchors that hold memory."
        ]
    if title == "Symbolism & Tattoos":
        return [
            "Symbols of continuity and navigation draw them—circles, maps, constellations, tides.",
            "Ink, if chosen, marks vows rather than decoration; milestones set into skin.",
            "Lines stay clean and intentional—the eye lands where intention lives."
        ]
    if title == "Motivation & Purpose":
        return [
            "Purpose grows from fidelity—show up long enough for meaning to mature.",
            "They are fueled by contribution over spectacle; craft over credit.",
            "Alignment compounds energy; the truer the path, the larger their risk budget."
        ]
    if title == "Communication & Social Energy":
        return [
            f"They {extra_p}; conversation needs contour and listening with weight.",
            "Tone is calibrated; truth without care feels like damage.",
            "If it can’t be said cleanly, it is not ready to be said."
        ]
    if title == "Relationship Blueprint":
        return [
            "Attachment is slow and intentional; safety is layered—consistency, honesty, steady affection.",
            "Boundaries are honored and asked in return; love is an ecosystem, not a performance.",
            "Betrayal doesn’t explode; it evaporates the future—access isn’t automatic after apology."
        ]
    if title == "Compatibility":
        return [
            "Best matches: partners who respect pace but add helpful novelty and honest dialogue.",
            "Stability plus curiosity amplifies their growth trajectory.",
            "Tough matches: volatile or performative styles that rush, test, or stage intimacy."
        ]
    if title == "Attraction & Seduction Style":
        return [
            "Attraction begins with safety and subtle magnetism—consistency over charisma.",
            "Seduction is attentiveness: remembered details, gentle humor, unfragmented presence.",
            "They fall for proof in motion—small acts that keep matching the words."
        ]
    if title == "Gifts, Taste & Interests":
        return [
            "They love artifacts with story—annotated books, well-made tools, photographs that hold a season.",
            "Experiences beat objects: a quiet travel day, a workshop, a designed surprise that respects bandwidth.",
            "Leisure blends restoration and creation—journaling, music, walking, craft."
        ]
    if title == "Career, Power & Productivity":
        return [
            "They excel where depth beats speed—research, design, therapy, editing, artisan craft, thoughtful product.",
            "Leadership is calm, principle-anchored, steady under pressure.",
            "Power accrues as reputation for delivery—not promises of delivery."
        ]
    if title == "Influence, Persuasion & Manipulation Profile":
        return [
            "They respond to clarity, consent, and patience—never to rush or scarcity theater.",
            "Risky levers *against* them: love-bombing, blurred boundaries, faux urgency when depleted.",
            "Defense: state pace, write definitions, sleep on big asks, verify instead of assume."
        ]
    if title == "Defense & Shadow":
        return [
            "Shadow shows as disengagement—quiet exits instead of clean confrontations.",
            "When hurt, they can become cool and procedural—truth kept, warmth lost.",
            "Integration work: keep the heart in the room while keeping standards intact."
        ]
    if title == "Growth Path & Transformation":
        return [
            "Weekly novelty that honors capacity: one new input, one micro-risk, one honest ask.",
            "Build rituals that metabolize emotion before it stacks—movement, sunlight, long-form writing.",
            "Practice real-time boundaries; prevention beats repair."
        ]
    if title == "Art, Entertainment & Culture Preferences":
        return [
            "They favor layered stories and music with motif; character arcs over spectacle.",
            "They revisit works that deepen on re-read—literary fiction, humanist cinema, ambient or piano-led music.",
            "Systems-aware games and emergent strategy fit their mind."
        ]
    if title == "Symbolic Archetype Summary":
        return [
            f"{name} moves like water—patient, shaping, resilient.",
            "A portable sanctuary follows them; rooms calm a few degrees when they enter.",
            "Their gift is continuity—keeping the thread when others forget why it mattered."
        ]
    if title == "MBTI Mirror":
        guess = "Closest MBTI varies by nuance"
        if (O,C,E,A,N)==("Low","Low","Low","Low","Low"): guess = "ISTP / ISFP"
        if (O,C,E,A,N)==("High","Low","High","High","Low"): guess = "ENFP / ENFJ"
        if (O,C,E,A,N)==("Low","High","Low","High","Low"): guess = "ISFJ / INFJ"
        return [
            f"Nearest mirrors by energy pattern: {guess}.",
            "MBTI and Big Five slice the psyche differently—treat this as a metaphor, not a box.",
            "Study mirrors for tactics, not identity swaps."
        ]
    if title == "Final Quote":
        return [ "“Move at the speed of truth. What remains at that pace is yours.”",
                 "“Let the work make the noise; you keep the signal.”",
                 "“Consistency is the quiet form of courage.”" ]
    # default fallback
    return [f"{name} deepens this domain.", "They refine rather than rush.", "Subtle changes compound."]

def build_doc():
    mapping = load_map()
    if not mapping:
        print("No archetype map found."); return
    doc = Document()
    doc.add_heading("Big Five Personality Archetypes — Full Tone A Edition", 0)

    # stable order
    for code in sorted(mapping.keys()):
        name = mapping[code]
        O,C,E,A,N = split5(code)
        doc.add_page_break()
        header = f"Openness: {O} | Conscientiousness: {C} | Extraversion: {E} | Agreeableness: {A} | Neuroticism: {N} — Archetype: {name}"
        h = doc.add_heading(header, level=1); h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for title in SECTION_TITLES:
            doc.add_heading(title, level=2)
            for para in three_paras(name, title, O,C,E,A,N):
                doc.add_paragraph(para)

    doc.save(OUT_DOC)
    print(f"✅ Created {OUT_DOC} with {len(mapping)} archetypes.")

if __name__ == "__main__":
    build_doc()