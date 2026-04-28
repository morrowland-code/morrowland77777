from flask import Flask, render_template, request, jsonify, send_file, redirect, session, abort
import os, re, json, secrets
from docx import Document
from io import BytesIO
from dotenv import load_dotenv
from flask import send_from_directory
import stripe

# ------------------------------------------------------------------
# 1️⃣ Config & Setup
# ------------------------------------------------------------------
load_dotenv()

app = Flask(__name__)

# Decide cookie security based on environment / DOMAIN
FLASK_ENV = os.getenv("FLASK_ENV", "development").lower()
DOMAIN = os.getenv("DOMAIN", "http://localhost:5000")

USE_SECURE_COOKIES = DOMAIN.startswith("https://") or FLASK_ENV == "production"

app.config.update(
    SESSION_COOKIE_SECURE=USE_SECURE_COOKIES,  # ✅ Secure in production / HTTPS, off on localhost
    SESSION_COOKIE_HTTPONLY=True,              # ✅ Not accessible via JavaScript
    SESSION_COOKIE_SAMESITE="Lax",             # ✅ Helps prevent CSRF
)

app.secret_key = os.getenv("FLASK_SECRET_KEY", secrets.token_hex(16))
stripe.api_key = os.getenv("STRIPE_SECRET_KEY", "")
OWNER_SECRET = os.getenv("OWNER_SECRET", "supersecureadminkey123")

print(f"[CONFIG] FLASK_ENV={FLASK_ENV}, DOMAIN={DOMAIN}, USE_SECURE_COOKIES={USE_SECURE_COOKIES}")

# ------------------------------------------------------------------
# 2️⃣ Parse morrowland 243.docx
# ------------------------------------------------------------------



def load_detailed_archetypes_text(file_path: str):
    if not os.path.exists(file_path):
        print(f"[ERROR] File not found: {file_path}")
        return {}, {}, {}

    with open(file_path, "r", encoding="utf-8") as f:
        raw = f.read()

    # Split on the marker
    blocks = [b.strip() for b in raw.split("====================") if b.strip()]

    by_code = {}
    by_name = {}
    code_to_name = {}
    duplicates = []

    header_re = re.compile(
        r"Openness\s*:\s*(Low|Medium|High)\s*\|\s*"
        r"Conscientiousness\s*:\s*(Low|Medium|High)\s*\|\s*"
        r"Extraversion\s*:\s*(Low|Medium|High)\s*\|\s*"
        r"Agreeableness\s*:\s*(Low|Medium|High)\s*\|\s*"
        r"Neuroticism\s*:\s*(Low|Medium|High)",
        re.IGNORECASE
    )

    archetype_re = re.compile(
        r"Archetype\s*:\s*(.+?)(?:\s*\(([HML]{5})\))?\s*$",
        re.IGNORECASE
    )

    for idx, block in enumerate(blocks, start=1):
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        if len(lines) < 2:
            print(f"[WARN] Skipping short block #{idx}")
            continue

        trait_line = None
        archetype_line = None

        for line in lines:
            if not trait_line and header_re.search(line):
                trait_line = line
            if not archetype_line and archetype_re.search(line):
                archetype_line = line

        if not trait_line:
            print(f"[WARN] Could not parse trait line in block #{idx}: {lines[:3]}")
            continue

        if not archetype_line:
            print(f"[WARN] Could not parse archetype line in block #{idx}: {lines[:3]}")
            continue

        m_header = header_re.search(trait_line)
        O, C, E, A, N_ = [x.capitalize() for x in m_header.groups()]
        code = f"{O}-{C}-{E}-{A}-{N_}"

        m_name = archetype_re.search(archetype_line)
        archetype_name = m_name.group(1).strip()

        # Keep the full block text exactly as report content
        report_text = "\n".join(lines).strip()

        if code in by_code:
            duplicates.append((code, code_to_name.get(code), archetype_name))
            print(f"[DUPLICATE CODE] {code}: '{code_to_name.get(code)}' overwritten by '{archetype_name}'")

        by_code[code] = report_text
        by_name[archetype_name] = report_text
        code_to_name[code] = archetype_name

    # Build all 243 possible codes
    levels = ["Low", "Medium", "High"]
    all_codes = {
        f"{o}-{c}-{e}-{a}-{n}"
        for o in levels for c in levels for e in levels for a in levels for n in levels
    }

    loaded_codes = set(by_code.keys())
    missing_codes = sorted(all_codes - loaded_codes)

    print(f"[✅ SUCCESS] Loaded {len(by_code)} archetypes from {os.path.basename(file_path)}")

    if duplicates:
        print(f"[⚠️ DUPLICATES FOUND: {len(duplicates)}]")
        for code, old_name, new_name in duplicates:
            print(f"   {code}: '{old_name}' -> '{new_name}'")

    if missing_codes:
        print(f"[⚠️ MISSING CODES: {len(missing_codes)}]")
        for code in missing_codes:
            print(f"   {code}")

    return by_code, by_name, code_to_name

# ------------------------------------------------------------------
# 3️⃣ Load Data
# ------------------------------------------------------------------
base_dir = os.path.dirname(os.path.abspath(__file__))
file_path = os.path.join(base_dir, "archetypes_cleaned.txt")

DETAILED_BY_CODE, DETAILED_BY_NAME, CODE_TO_NAME = load_detailed_archetypes_text(file_path)


def load_archetypes(base_map):
    """
    Build the archetype-name mapping.
    1. Start with names parsed from the DOCX (CODE_TO_NAME).
    2. If archetypes_full.json / archetypes.json exist, they override or add.
    """
    result = dict(base_map or {})
    for file in ["archetypes_full.json", "archetypes.json"]:
        path = os.path.join(base_dir, file)
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
                if isinstance(data, dict) and data:
                    result.update(data)
                    print(f"[INFO] Loaded {len(data)} archetypes from {file}")
                    return result
    if not result:
        print("[⚠️ Using fallback minimal archetypes]")
        return {"Low-Low-Low-Low-Low": "Aquashine"}
    return result


ARCHETYPES = load_archetypes(CODE_TO_NAME)
FREE_CODES_FILE = os.path.join(base_dir, "free_codes.json")

# ------------------------------------------------------------------
# 4️⃣ Free Code System
# ------------------------------------------------------------------
def load_free_codes():
    if os.path.exists(FREE_CODES_FILE):
        try:
            with open(FREE_CODES_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except json.JSONDecodeError:
            print("[⚠️ Corrupted free_codes.json, resetting.]")
            return {}
    return {}


def save_free_codes(codes):
    with open(FREE_CODES_FILE, "w", encoding="utf-8") as f:
        json.dump(codes, f, indent=2)


def generate_free_code():
    code = secrets.token_hex(4).upper()
    codes = load_free_codes()
    codes[code] = {"used": False}
    save_free_codes(codes)
    print(f"[🎁 NEW CODE GENERATED]: {code}")
    return code


def verify_free_code(code: str, owner_key: str = ""):
    if owner_key == OWNER_SECRET:
        print(f"[👑 OWNER BYPASS CODE ACCEPTED]: {code}")
        return True

    codes = load_free_codes()
    if code in codes and not codes[code]["used"]:
        codes[code]["used"] = True
        save_free_codes(codes)
        print(f"[✅ FREE CODE ACCEPTED]: {code}")
        return True

    print(f"[❌ INVALID/USED CODE]: {code}")
    return False

# ------------------------------------------------------------------
# 5️⃣ Socials
# ------------------------------------------------------------------
@app.context_processor
def inject_socials():
    return dict(
        tiktok_url="https://www.tiktok.com/@neptunee7777",
        instagram_url="https://www.instagram.com/kendallm16",
    )


@app.route('/robots.txt')
def robots():
    return send_from_directory('static', 'robots.txt')



# ------------------------------------------------------------------
# 6️⃣ Routes
# ------------------------------------------------------------------
@app.route("/")
def index():
    return render_template("index.html")


@app.route("/quiz")
def quiz():
    return render_template("quiz.html")


# 🔐 Owner-only: Generate free codes
@app.route("/generate-free-code")
def make_free_code():
    key = request.args.get("key", "")
    if key != OWNER_SECRET:
        abort(403)
    code = generate_free_code()
    return jsonify({"new_code": code})


@app.route("/verify-free-code", methods=["POST"])
def api_verify_free_code():
    data = request.get_json() or {}
    code = (data.get("code") or "").strip().upper()
    owner_key = (data.get("owner_key") or "").strip()

    if not code:
        return jsonify({"valid": False, "error": "No code provided."}), 400

    if verify_free_code(code, owner_key=owner_key):
        code = session.get("latest_code")

        unlocked = session.get("unlocked_codes", [])

        if code and code not in unlocked:
            unlocked.append(code)

        session["unlocked_codes"] = unlocked

        return jsonify({"valid": True})

    return jsonify({"valid": False, "error": "Invalid or already used."}), 400

# 🧠 Save quiz code to session (from frontend)
@app.route("/api/set-latest-code", methods=["POST"])
def set_latest_code():
    data = request.get_json() or {}
    code = data.get("code", "")
    if not code:
        return jsonify({"success": False, "error": "No code provided"}), 400
    session["latest_code"] = code
    print(f"[🧠 SAVED QUIZ CODE]: {code}")
    return jsonify({"success": True})


# 💳 Stripe Checkout
@app.route("/create-checkout-session")
def create_checkout_session():
    """
    We assume /api/set-latest-code has already stored the user's code.
    Do NOT overwrite latest_code here.
    """
    user_id = secrets.token_hex(6)
    checkout_token = secrets.token_urlsafe(16)

    session["user_id"] = user_id
    session["checkout_token"] = checkout_token  # secret random value for this flow

    try:
        stripe_session = stripe.checkout.Session.create(
            mode="payment",
            payment_method_types=["card"],
            line_items=[{
                "price_data": {
                    "currency": "usd",
                    "product_data": {"name": "Big 5 Detailed Archetype Report"},
                    "unit_amount": 99,
                },
                "quantity": 1
            }],
            # include both Stripe session_id and our own secret token
            success_url=f"{DOMAIN}/purchase-success?session_id={{CHECKOUT_SESSION_ID}}&token={checkout_token}",
            cancel_url=f"{DOMAIN}/",
        )

        # store Stripe's session id in our Flask session for verification
        session["checkout_session_id"] = stripe_session.id

        return redirect(stripe_session.url)
    except Exception as e:
        print("Stripe error:", e)
        return f"Stripe session creation failed: {e}", 500

@app.route("/get-archetype-by-name")
def get_by_name():
    name = request.args.get("name")
    return DETAILED_BY_NAME.get(name, "Not found")

@app.route("/browse-by-traits")
def browse_by_traits():
    return render_template("browse_by_traits.html")

@app.route("/select-by-traits", methods=["POST"])
def select_by_traits():
    O = (request.form.get("openness") or "").strip()
    C = (request.form.get("conscientiousness") or "").strip()
    E = (request.form.get("extraversion") or "").strip()
    A = (request.form.get("agreeableness") or "").strip()
    N = (request.form.get("neuroticism") or "").strip()

    key = f"{O}-{C}-{E}-{A}-{N}"

    if key not in DETAILED_BY_CODE:
        return f"Archetype not found for {key}", 404

    session["latest_code"] = key

    unlocked = session.get("unlocked_codes", [])
    if key in unlocked:
        return redirect("/report")

    return redirect("/unlock-report")

@app.route('/sitemap.xml')
def sitemap():
    return send_from_directory('static', 'sitemap.xml')


@app.route("/unlock-report")
def unlock_report():
    code = session.get("latest_code")
    if not code:
        return redirect("/browse")

    archetype_name = ARCHETYPES.get(code, "Unknown Archetype")
    return render_template("unlock_report.html", archetype=archetype_name, traits=code)



@app.route("/test-archetype/<code>")
def test_archetype(code):
    result = DETAILED_BY_CODE.get(code)

    if not result:
        return "Not found"

    return f"""
    <h1>{CODE_TO_NAME.get(code, 'Unknown')}</h1>
    <h2>{code}</h2>
    <pre style="white-space: pre-wrap;">{result}</pre>
    """


@app.route("/get-archetype", methods=["POST"])
def get_archetype():
    data = request.json

    O = data.get("openness")
    C = data.get("conscientiousness")
    E = data.get("extraversion")
    A = data.get("agreeableness")
    N = data.get("neuroticism")

    key = f"{O}-{C}-{E}-{A}-{N}"

    result = DETAILED_BY_CODE.get(key)

    if not result:
        return {"error": "Archetype not found"}, 404

    return {
        "code": key,
        "archetype": CODE_TO_NAME.get(key),
        "description": result
    }


@app.route("/purchase-success")
def purchase_success():
    session_id = request.args.get("session_id", "")
    token = request.args.get("token", "")

    expected_id = session.get("checkout_session_id")
    expected_token = session.get("checkout_token")

    if not session_id or not token or session_id != expected_id or token != expected_token:
        print("[SECURITY] Invalid purchase-success attempt.")
        abort(403)

    try:
        stripe_session = stripe.checkout.Session.retrieve(session_id)
        if stripe_session.payment_status != "paid":
            print("[SECURITY] Stripe session not paid.")
            abort(403)
    except Exception as e:
        print("[Stripe verify error]", e)
        abort(403)

    code = session.get("latest_code")
    unlocked = session.get("unlocked_codes", [])

    if code and code not in unlocked:
        unlocked.append(code)

    session["unlocked_codes"] = unlocked
    session.pop("checkout_session_id", None)
    session.pop("checkout_token", None)

    return redirect("/report")

@app.route("/debug-reset-session")
def debug_reset_session():
    session.clear()
    return "Session cleared."



# 🔒 Secure report route (with archetype name)
@app.route("/report")
def report():
    code = session.get("latest_code")
    unlocked = session.get("unlocked_codes", [])

    if not code or code not in unlocked:
        return redirect("/unlock-report")

    code = session.get("latest_code", "Medium-Medium-Medium-Medium-Medium")
    archetype_name = ARCHETYPES.get(code, None)
    detailed_text = DETAILED_BY_CODE.get(code)

    if not detailed_text and archetype_name:
        detailed_text = DETAILED_BY_NAME.get(archetype_name)

    if not detailed_text:
        detailed_text = "Detailed report not found."

    if not archetype_name:
        archetype_name = "Unknown Archetype"

    return render_template(
        "detailed_report.html",
        archetype=archetype_name,
        traits=code,
        sections={"Detailed Report": detailed_text},
    )


@app.route("/api/render-report")
def api_render_report():
    code = session.get("latest_code")
    unlocked = session.get("unlocked_codes", [])

    if not code or code not in unlocked:
        abort(403)

    code = session.get("latest_code", "Medium-Medium-Medium-Medium-Medium")
    archetype_name = ARCHETYPES.get(code, None)
    detailed_text = DETAILED_BY_CODE.get(code)

    if not detailed_text and archetype_name:
        detailed_text = DETAILED_BY_NAME.get(archetype_name)

    if not detailed_text:
        detailed_text = "Detailed report not found."

    if not archetype_name:
        archetype_name = "Unknown Archetype"

    return render_template(
        "detailed_report.html",
        archetype=archetype_name,
        traits=code,
        sections={"Detailed Report": detailed_text},
        quote="“Depth rewards patience.”",
    )

@app.route("/browse")
def browse():
    return render_template("browse.html", ARCHETYPES=ARCHETYPES)


@app.route("/select-archetype", methods=["POST"])
def select_archetype():
    code = (request.form.get("code") or "").strip()

    if not code or code not in ARCHETYPES:
        return redirect("/browse")

    session["latest_code"] = code

    unlocked = session.get("unlocked_codes", [])
    if code in unlocked:
        return redirect("/report")

    return redirect("/unlock-report")


@app.route("/reset-free-code")
def reset_free_code():
    key = request.args.get("key", "")
    code = (request.args.get("code", "") or "").strip().upper()

    if key != OWNER_SECRET:
        abort(403)

    codes = load_free_codes()

    if code not in codes:
        return jsonify({"success": False, "error": "Code not found"}), 404

    codes[code]["used"] = False
    save_free_codes(codes)

    return jsonify({"success": True, "code": code, "used": False})

@app.route("/api/download-report")
def download_report():
    code = session.get("latest_code") or request.args.get("code", "")
    unlocked = session.get("unlocked_codes", [])

    if not code or code not in unlocked:
        abort(403)

    name = ARCHETYPES.get(code, "Unknown")
    detailed_text = DETAILED_BY_CODE.get(code) or DETAILED_BY_NAME.get(name)

    doc = Document()
    doc.add_heading(name, level=1)
    doc.add_paragraph(detailed_text or "Detailed text not found.")

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return send_file(
        output,
        as_attachment=True,
        download_name=f"{name.replace(' ', '_')}_Detailed_Report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

# ------------------------------------------------------------------
# 🏁 Run Flask
# ------------------------------------------------------------------
if __name__ == "__main__":
    # debug=True is fine for local; in production you’ll run via gunicorn/https
    app.run(debug=False)
